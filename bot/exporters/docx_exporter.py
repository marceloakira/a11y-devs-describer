import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from bot.utils.logger import logger
from bot.utils.text_processor import parse_markdown_and_descriptions

def export_docx(text: str, output_path: Path, filename: str = "") -> Path:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)

    if filename:
        heading = doc.add_heading(filename, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    structured_content = parse_markdown_and_descriptions(text)
    
    for entry_type, content in structured_content:
        if entry_type == 'h1':
            doc.add_heading(content, level=1)
        elif entry_type == 'h2':
            doc.add_heading(content, level=2)
        elif entry_type == 'h3':
            doc.add_heading(content, level=3)
        elif entry_type == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, content)
        elif entry_type == 'number':
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, content)
        elif entry_type == 'description':
            # No Word, colocamos a descrição em um box destacado para acessibilidade
            p = doc.add_paragraph()
            run = p.add_run("[INÍCIO DA AUDIODESCRIÇÃO]")
            run.bold = True
            run.font.size = Pt(10)
            
            desc_para = doc.add_paragraph(content)
            desc_para.paragraph_format.left_indent = Cm(1.0)
            for run in desc_para.runs: 
                run.font.size = Pt(11)
                run.italic = True
                
            p_end = doc.add_paragraph()
            run_end = p_end.add_run("[FIM DA AUDIODESCRIÇÃO]")
            run_end.bold = True
            run_end.font.size = Pt(10)
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, content)
            p.space_after = Pt(6)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.debug("DOCX exportado: {}", output_path)
    return output_path

def _add_formatted_text(paragraph, text):
    """Analisa Markdown básico (**bold**, *italic*) e adiciona ao parágrafo do Word."""
    # Regex para capturar negrito, itálico e ambos
    parts = re.split(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)", text)
    
    for part in parts:
        if not part: continue
        
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def _collect_table_lines(lines: list[str]) -> list[list[str]] | None:
    table_data: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            if all(c in "| -:" for c in stripped) and "---" in stripped:
                continue
            cells = [c.strip() for c in stripped.split("|") if c.strip()]
            if len(cells) >= 2:
                table_data.append(cells)
    return table_data if table_data else None

def _add_table_to_doc(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    doc.add_paragraph()
